from datetime import datetime
from database.session import SessionLocal
from database.models import Event, GitActivity, Report
import os
from jinja2 import Template
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors


def summarize_events(start: datetime, end: datetime) -> str:
    session = SessionLocal()
    events = session.query(Event).filter(
        Event.timestamp >= start, Event.timestamp <= end
    ).order_by(Event.timestamp).all()

    git_activity = session.query(GitActivity).filter(
        GitActivity.timestamp >= start, GitActivity.timestamp <= end
    ).all()
    session.close()

    lines = []
    lines.append(f"Report: {start.strftime('%Y-%m-%d %H:%M')} → {end.strftime('%Y-%m-%d %H:%M')}")
    lines.append("=" * 55)

    # --- Category breakdown ---
    cat_map = {}
    for e in events:
        if e.idle:
            continue
        cat = e.category or 'Other'
        cat_map.setdefault(cat, 0)
        cat_map[cat] += e.duration or 0

    total_secs = sum(cat_map.values()) or 1
    lines.append("\n🗂  ACTIVITY CATEGORIES")
    lines.append("-" * 30)
    for cat, secs in sorted(cat_map.items(), key=lambda x: -x[1]):
        pct = secs / total_secs * 100
        hours = secs / 3600.0
        bar = '█' * int(pct / 5)
        lines.append(f"  {cat:<22} {hours:5.2f}h  ({pct:4.1f}%)  {bar}")

    # --- Browser History (from BrowserHistory table) ---
    browser_lines = _get_browser_history_section(start, end)
    if browser_lines:
        lines.extend(browser_lines)

    # --- Website breakdown (from Events table — fallback) ---
    web_map = {}
    for e in events:
        if e.idle or not e.website:
            continue
        web_map.setdefault(e.website, 0)
        web_map[e.website] += e.duration or 0

    if web_map:
        lines.append("\n🌐  WEBSITES & APPS VISITED")
        lines.append("-" * 30)
        for site, secs in sorted(web_map.items(), key=lambda x: -x[1])[:15]:
            mins = secs / 60.0
            lines.append(f"  {site:<30} {mins:5.1f} min")

    # --- Files edited (enhanced with IDE info) ---
    file_lines = _get_file_edits_section(start, end)
    if file_lines:
        lines.extend(file_lines)
    else:
        # Fallback to basic file tracking from Events
        file_map = {}
        for e in events:
            if e.opened_file:
                file_map.setdefault(e.opened_file, 0)
                file_map[e.opened_file] += e.duration or 0

        if file_map:
            lines.append("\n💻  FILES WORKED ON")
            lines.append("-" * 30)
            for fname, secs in sorted(file_map.items(), key=lambda x: -x[1])[:10]:
                mins = secs / 60.0
                lines.append(f"  {fname:<40} {mins:4.1f} min")

    # --- Git commits ---
    if git_activity:
        lines.append("\n🔀  GIT COMMITS")
        lines.append("-" * 30)
        for ga in git_activity[:10]:
            msg = (ga.message or '').strip().split('\n')[0][:60]
            lines.append(f"  [{ga.commit_hash[:7]}] {msg}")
            lines.append(f"           by {ga.author}  @ {ga.timestamp.strftime('%H:%M')}")

    lines.append("\n" + "=" * 55)
    lines.append(f"Total tracked time: {total_secs/3600:.2f} hours")

    return "\n".join(lines)


def _get_browser_history_section(start: datetime, end: datetime) -> list[str]:
    """Generate the Browser History section from the BrowserHistory table."""
    lines = []
    try:
        from database.models import BrowserHistory
        session = SessionLocal()
        rows = (session.query(BrowserHistory)
                .filter(BrowserHistory.timestamp >= start,
                        BrowserHistory.timestamp <= end)
                .order_by(BrowserHistory.timestamp.desc())
                .limit(50).all())
        session.close()

        if not rows:
            return []

        # Aggregate by domain
        domain_map = {}  # domain → {site_name, visits, category, urls}
        for r in rows:
            d = r.domain or 'Unknown'
            if d not in domain_map:
                domain_map[d] = {
                    'site_name': r.site_name or d,
                    'visits': 0,
                    'category': r.category or 'Browsing',
                    'urls': [],
                }
            domain_map[d]['visits'] += 1
            if r.url and len(domain_map[d]['urls']) < 3:
                title = (r.title or r.url)[:50]
                domain_map[d]['urls'].append({
                    'title': title,
                    'url': r.url[:80],
                    'time': r.timestamp.strftime('%H:%M') if r.timestamp else '',
                })

        lines.append("\n🔍  BROWSER HISTORY")
        lines.append("-" * 70)
        lines.append(f"  {'Site':<25} {'Visits':>6}  {'Category':<20}")
        lines.append(f"  {'─' * 25} {'─' * 6}  {'─' * 20}")

        for domain, info in sorted(domain_map.items(), key=lambda x: -x[1]['visits'])[:20]:
            name = info['site_name'][:25]
            lines.append(f"  {name:<25} {info['visits']:>6}  {info['category']:<20}")
            # Show up to 3 pages visited
            for u in info['urls']:
                title = u['title'][:45]
                lines.append(f"    └ [{u['time']}] {title}")
                if u['url'] != u['title']:
                    url_short = u['url'][:65]
                    lines.append(f"      {url_short}")

    except Exception as e:
        pass  # BrowserHistory table may not exist on older DBs
    return lines


def _get_file_edits_section(start: datetime, end: datetime) -> list[str]:
    """Generate the Files Edited section from the FileEdit table with IDE info."""
    lines = []
    try:
        from database.models import FileEdit
        session = SessionLocal()
        today = start.strftime('%Y-%m-%d')
        rows = (session.query(FileEdit)
                .filter(FileEdit.timestamp >= start,
                        FileEdit.timestamp <= end)
                .order_by(FileEdit.duration_sec.desc())
                .limit(20).all())
        session.close()

        if not rows:
            return []

        lines.append("\n💻  FILES EDITED IN CODING")
        lines.append("-" * 80)
        lines.append(f"  {'File':<30} {'IDE':<14} {'Project':<16} {'Lang':<10} {'Time':>8}")
        lines.append(f"  {'─' * 30} {'─' * 14} {'─' * 16} {'─' * 10} {'─' * 8}")

        for r in rows:
            fname = (r.file_name or r.file_path or 'unknown')[:30]
            editor = (r.editor or '-')[:14]
            project = (r.project or '-')[:16]
            lang = (r.language or '-')[:10]
            mins = (r.duration_sec or 0) / 60.0
            time_str = f"{mins:5.1f} min"
            lines.append(f"  {fname:<30} {editor:<14} {project:<16} {lang:<10} {time_str:>8}")

    except Exception:
        pass  # FileEdit table may not exist on older DBs
    return lines


def generate_markdown(start: datetime, end: datetime, out_folder: str):
    summary = summarize_events(start, end)
    fname = os.path.join(out_folder, f"report_{start.date()}_{end.date()}.md")
    with open(fname, "w", encoding="utf-8") as f:
        f.write("# WorkSense AI — Daily Activity Report\n\n")
        f.write(f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n\n")
        f.write("```\n")
        f.write(summary)
        f.write("\n```\n")
    return fname


def generate_docx(start: datetime, end: datetime, out_folder: str):
    summary = summarize_events(start, end)
    doc = Document()

    # Title
    title = doc.add_heading('WorkSense AI — Activity Report', level=1)
    title.runs[0].font.color.rgb = RGBColor(0x1E, 0x90, 0xFF)

    doc.add_paragraph(f"Period: {start.strftime('%Y-%m-%d %H:%M')} → {end.strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph()

    for line in summary.split('\n'):
        p = doc.add_paragraph(line)
        if line.startswith('🗂') or line.startswith('🌐') or line.startswith('💻') or line.startswith('🔀') or line.startswith('🔍'):
            p.runs[0].bold = True if p.runs else False

    fname = os.path.join(out_folder, f"report_{start.date()}_{end.date()}.docx")
    doc.save(fname)
    return fname


def generate_pdf(start: datetime, end: datetime, out_folder: str):
    summary = summarize_events(start, end)
    fname = os.path.join(out_folder, f"report_{start.date()}_{end.date()}.pdf")
    c = canvas.Canvas(fname, pagesize=letter)

    # Header
    c.setFillColor(colors.HexColor('#1E90FF'))
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, 770, "WorkSense AI — Activity Report")
    c.setFillColor(colors.black)
    c.setFont("Helvetica", 9)
    c.drawString(50, 752, f"Period: {start.strftime('%Y-%m-%d %H:%M')} to {end.strftime('%Y-%m-%d %H:%M')}")
    c.drawString(50, 740, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    y = 720
    c.setFont("Courier", 8)
    for line in summary.splitlines():
        if y < 50:
            c.showPage()
            y = 770
            c.setFont("Courier", 8)
        # Highlight section headers
        section_markers = ('🗂', '🌐', '💻', '🔀', '🔍')
        if any(line.startswith(m) for m in section_markers):
            c.setFont("Helvetica-Bold", 9)
            c.setFillColor(colors.HexColor('#1E90FF'))
        else:
            c.setFont("Courier", 8)
            c.setFillColor(colors.black)
        c.drawString(50, y, line.encode('ascii', 'replace').decode())
        y -= 12

    c.save()
    return fname
